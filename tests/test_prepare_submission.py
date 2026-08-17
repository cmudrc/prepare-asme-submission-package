from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches
from lxml import etree
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


REPO = Path(__file__).resolve().parents[1]
SCRIPT = REPO / "scripts" / "prepare_submission.py"


def linework(path: Path, pixels: tuple[int, int], dpi: int = 900) -> None:
    image = Image.new("RGB", pixels, "white")
    draw = ImageDraw.Draw(image)
    margin = max(10, pixels[0] // 30)
    draw.rectangle((margin, margin, pixels[0] - margin, pixels[1] - margin), outline="black", width=max(3, pixels[0] // 400))
    draw.line((margin, pixels[1] - margin, pixels[0] // 2, margin, pixels[0] - margin, pixels[1] - margin), fill="black", width=max(3, pixels[0] // 450))
    draw.text((margin * 2, margin * 2), "DESIGN SPACE", fill="black")
    image.save(path, dpi=(dpi, dpi))


def word_fixture(path: Path, image_path: Path) -> None:
    document = Document()
    document.add_heading("A Validated Design Manuscript", 0)
    document.add_paragraph("A. Author, Carnegie Mellon University, author@example.edu")
    document.add_heading("Abstract", level=1)
    document.add_paragraph("This fixture validates final-production packaging.")
    document.add_paragraph("Figure 1 shows the feasible region and remains legible at final size.")
    document.add_picture(str(image_path), width=Inches(3.0))
    document.add_paragraph("Fig. 1. A linework design-space diagram.")
    document.add_heading("Data Availability", level=1)
    document.add_paragraph("The fixture data are available at https://example.org/data.")
    document.add_heading("Conflict of Interest", level=1)
    document.add_paragraph("The authors declare no conflicts of interest.")
    document.add_heading("References", level=1)
    document.add_paragraph("[1] A. Author, 2026, A Fixture Reference.")
    document.save(path)


def inject_word_format_revisions(path: Path) -> None:
    """Add synthetic property changes and a fully deleted paragraph."""
    replacement = path.with_suffix(".revised.docx")
    with zipfile.ZipFile(path) as source, zipfile.ZipFile(replacement, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "word/document.xml":
                namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                qname = lambda value: f"{{{namespace}}}{value}"
                root = etree.fromstring(data)
                paragraph = root.find(f".//{qname('p')}")
                run = root.find(f".//{qname('r')}")
                assert paragraph is not None and run is not None
                p_properties = paragraph.find(qname("pPr"))
                if p_properties is None:
                    p_properties = etree.Element(qname("pPr"))
                    paragraph.insert(0, p_properties)
                r_properties = run.find(qname("rPr"))
                if r_properties is None:
                    r_properties = etree.Element(qname("rPr"))
                    run.insert(0, r_properties)
                p_change = etree.SubElement(
                    p_properties,
                    qname("pPrChange"),
                    {qname("id"): "100", qname("author"): "Fixture"},
                )
                etree.SubElement(p_change, qname("pPr"))
                r_change = etree.SubElement(
                    r_properties,
                    qname("rPrChange"),
                    {qname("id"): "101", qname("author"): "Fixture"},
                )
                etree.SubElement(r_change, qname("rPr"))
                body = root.find(f".//{qname('body')}")
                assert body is not None
                deleted_paragraph = etree.Element(qname("p"))
                deleted = etree.SubElement(
                    deleted_paragraph,
                    qname("del"),
                    {qname("id"): "102", qname("author"): "Fixture"},
                )
                deleted_run = etree.SubElement(deleted, qname("r"))
                deleted_text = etree.SubElement(deleted_run, qname("delText"))
                deleted_text.text = "Deleted fixture paragraph"
                section = body.find(qname("sectPr"))
                body.insert(body.index(section) if section is not None else len(body), deleted_paragraph)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            target.writestr(item, data)
    replacement.replace(path)


def word_fixture_with_inline_graphic(path: Path, figure_path: Path, inline_path: Path) -> None:
    document = Document()
    document.add_heading("A Validated Design Manuscript", 0)
    document.add_paragraph("A. Author, Carnegie Mellon University, author@example.edu")
    document.add_heading("Abstract", level=1)
    document.add_paragraph("An inline notation symbol appears in the table and is not a numbered figure.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Notation"
    table.cell(0, 1).paragraphs[0].add_run().add_picture(str(inline_path), width=Inches(0.4))
    document.add_paragraph("The feasible region is shown in Fig. 1.")
    document.add_picture(str(figure_path), width=Inches(3.0))
    document.add_paragraph("Fig. 1. A linework design-space diagram.")
    document.add_heading("Data Availability", level=1)
    document.add_paragraph("The fixture data are available at https://example.org/data.")
    document.add_heading("Conflict of Interest", level=1)
    document.add_paragraph("The authors declare no conflicts of interest.")
    document.save(path)


def word_fixture_with_grouped_callout(path: Path, image_path: Path) -> None:
    document = Document()
    document.add_heading("A Validated Design Manuscript", 0)
    document.add_paragraph("Figures 1 and 2 compare the baseline and revised design spaces.")
    document.add_picture(str(image_path), width=Inches(3.0))
    document.add_paragraph("Fig. 1. The baseline design space.")
    document.add_picture(str(image_path), width=Inches(3.0))
    document.add_paragraph("Fig. 2. The revised design space.")
    document.add_heading("Data Availability", level=1)
    document.add_paragraph("The fixture data are available at https://example.org/data.")
    document.add_heading("Conflict of Interest", level=1)
    document.add_paragraph("The authors declare no conflicts of interest.")
    document.save(path)


def complete_pdf(path: Path, image_path: Path) -> None:
    pdf = canvas.Canvas(str(path), pagesize=letter)
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(72, 740, "A Validated LaTeX Manuscript")
    pdf.setFont("Helvetica", 10)
    pdf.drawString(72, 718, "The feasible region is shown in Fig. 1 [1].")
    pdf.drawImage(str(image_path), 72, 480, width=216, height=108, preserveAspectRatio=True)
    pdf.drawString(72, 465, "Fig. 1. A linework design-space diagram.")
    pdf.drawString(72, 440, "Data Availability: https://example.org/data")
    pdf.drawString(72, 425, "Conflict of Interest: None.")
    pdf.drawString(72, 390, "References")
    pdf.drawString(72, 374, "[1] A. Author, A Fixture Reference, 2026.")
    pdf.save()


def complete_pdf_with_reference_traps(path: Path, image_path: Path) -> None:
    pdf = canvas.Canvas(str(path), pagesize=letter)
    pdf.setFont("Helvetica", 10)
    pdf.drawString(72, 740, "A Validated LaTeX Manuscript")
    pdf.drawImage(str(image_path), 72, 500, width=216, height=108, preserveAspectRatio=True)
    pdf.drawString(72, 470, "References")
    for index in range(1, 6):
        suffix = "References in Design" if index == 5 else "A Fixture Reference"
        pdf.drawString(72, 470 - index * 16, f"[{index}] A. Author, {suffix}.")
    pdf.drawString(72, 370, "2002 AAAI Spring Symposium")
    pdf.drawString(72, 354, "2003 Another Proceedings")
    pdf.drawString(72, 338, "2005 Final Proceedings")
    pdf.save()


def complete_pdf_with_blank_page(path: Path, image_path: Path) -> None:
    complete_pdf(path, image_path)
    original = path.with_suffix(".content.pdf")
    path.replace(original)
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(str(original))
    writer = PdfWriter()
    writer.append_pages_from_reader(reader)
    writer.add_blank_page(width=letter[0], height=letter[1])
    with path.open("wb") as handle:
        writer.write(handle)
    original.unlink()


def latex_fixture(path: Path, image_path: Path, nocite: bool = False) -> None:
    main = r"""\documentclass{article}
\usepackage{graphicx}
\begin{document}
\title{A Validated LaTeX Manuscript}\maketitle
The feasible region in Fig.~\ref{fig:space} follows prior work~\cite{fixture}.
\begin{figure}
\centering
\includegraphics[width=3in]{figure1.png}
\caption{A linework design-space diagram.}
\label{fig:space}
\end{figure}
\section*{Data Availability}
The fixture data are available at \texttt{https://example.org/data}.
\section*{Conflict of Interest}
The authors declare no conflicts of interest.
%NOCITE%
\bibliographystyle{plain}
\bibliography{references}
\end{document}
"""
    main = main.replace("%NOCITE%", r"\nocite{*}" if nocite else "")
    bib = r"""@article{fixture,
  author = {Author, A.},
  title = {A Fixture Reference},
  year = {2026},
  journal = {Journal of Fixtures}
}
"""
    with tempfile.TemporaryDirectory() as temp:
        root = Path(temp)
        (root / "main.tex").write_text(main, encoding="utf-8")
        (root / "references.bib").write_text(bib, encoding="utf-8")
        shutil.copy2(image_path, root / "figure1.png")
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
            for source in sorted(root.iterdir()):
                archive.write(source, source.name)


class PrepareSubmissionTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory(prefix="jmd-skill-test-")
        self.root = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def run_packager(self, *arguments: str) -> subprocess.CompletedProcess[str]:
        environment = os.environ.copy()
        override = Path.home() / ".cache/codex-runtimes/codex-primary-runtime/dependencies/bin/override"
        environment["PATH"] = f"{override}:{environment.get('PATH', '')}"
        return subprocess.run(
            [sys.executable, str(SCRIPT), *arguments],
            text=True,
            capture_output=True,
            env=environment,
            check=False,
            timeout=90,
        )

    def test_word_package_preserves_text_and_extracts_figure(self) -> None:
        image = self.root / "figure.png"
        manuscript = self.root / "manuscript.docx"
        output = self.root / "word-output"
        linework(image, (2700, 1350))
        word_fixture(manuscript, image)
        inject_word_format_revisions(manuscript)

        result = self.run_packager(
            str(manuscript),
            "--paper-id", "MD-26-1001",
            "--journal", "jmd",
            "--output", str(output),
            "--require-section", "Data Availability",
            "--require-section", "Conflict of Interest",
        )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        self.assertTrue((output / "MD-26-1001_submission_package.zip").is_file())
        self.assertTrue((output / "upload/MD-26-1001_complete.pdf").is_file())
        self.assertTrue((output / "upload/figures/Fig_1.tiff").is_file())
        text_only = output / "upload/MD-26-1001_text_only.docx"
        with zipfile.ZipFile(text_only) as archive:
            self.assertFalse(any(name.startswith("word/media/") for name in archive.namelist()))
            xml = archive.read("word/document.xml").decode("utf-8")
            self.assertIn("Figure 1 shows", xml)
            self.assertIn("Fig. 1. A linework", xml)
        self.assertNotIn("UNCALLED_FIGURE", (output / "QA_REPORT.md").read_text())
        report = (output / "QA_REPORT.md").read_text()
        self.assertIn("removed 2 formatting-revision record(s)", report)
        self.assertIn("1 emptied deleted paragraph(s)", report)
        with zipfile.ZipFile(output / "work/MD-26-1001_clean.docx") as archive:
            clean_xml = archive.read("word/document.xml").decode("utf-8")
            self.assertNotIn("pPrChange", clean_xml)
            self.assertNotIn("rPrChange", clean_xml)
            self.assertNotIn("Deleted fixture paragraph", clean_xml)
        manifest = json.loads((output / "manifest.json").read_text())
        self.assertEqual(manifest["status"], "pass")
        self.assertTrue(manifest["manual_review_required"])
        self.assertGreaterEqual(manifest["figures"][0]["effective_dpi"], 899.9)
        self.assertTrue(list((output / "qa/complete").glob("*.png")))
        self.assertTrue(list((output / "qa/text-only").glob("*.png")))

    def test_low_effective_dpi_blocks_zip(self) -> None:
        image = self.root / "figure.png"
        manuscript = self.root / "manuscript.docx"
        output = self.root / "low-dpi-output"
        linework(image, (900, 450), dpi=300)
        word_fixture(manuscript, image)

        result = self.run_packager(
            str(manuscript),
            "--paper-id", "JCISE-26-1002",
            "--journal", "jcise",
            "--output", str(output),
        )

        self.assertEqual(result.returncode, 2, result.stdout + result.stderr)
        self.assertFalse((output / "JCISE-26-1002_submission_package.zip").exists())
        report = (output / "QA_REPORT.md").read_text()
        self.assertIn("LOW_EFFECTIVE_DPI", report)

    def test_word_text_only_preserves_unnumbered_semantic_graphic(self) -> None:
        figure = self.root / "figure.png"
        inline = self.root / "inline.png"
        manuscript = self.root / "manuscript-with-inline.docx"
        output = self.root / "word-inline-output"
        linework(figure, (2700, 1350))
        linework(inline, (240, 120), dpi=300)
        word_fixture_with_inline_graphic(manuscript, figure, inline)

        result = self.run_packager(
            str(manuscript),
            "--paper-id", "MD-26-1005",
            "--journal", "jmd",
            "--output", str(output),
        )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        text_only = output / "upload/MD-26-1005_text_only.docx"
        with zipfile.ZipFile(text_only) as archive:
            media = [name for name in archive.namelist() if name.startswith("word/media/")]
            self.assertEqual(len(media), 1)
        report = (output / "QA_REPORT.md").read_text()
        self.assertIn("INLINE_GRAPHICS_PRESERVED", report)
        self.assertNotIn("FIGURE_WITHOUT_CAPTION", report)

    def test_word_grouped_callout_credits_every_figure(self) -> None:
        figure = self.root / "figure.png"
        manuscript = self.root / "manuscript-with-grouped-callout.docx"
        output = self.root / "word-grouped-output"
        linework(figure, (2700, 1350))
        word_fixture_with_grouped_callout(manuscript, figure)

        result = self.run_packager(
            str(manuscript),
            "--paper-id", "MD-26-1008",
            "--journal", "jmd",
            "--output", str(output),
        )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        report = (output / "QA_REPORT.md").read_text()
        self.assertNotIn("UNCALLED_FIGURE", report)
        self.assertTrue((output / "upload/figures/Fig_1.tiff").is_file())
        self.assertTrue((output / "upload/figures/Fig_2.tiff").is_file())

    def test_overleaf_package_keeps_bib_separate_and_disables_graphics(self) -> None:
        image = self.root / "figure1.png"
        source = self.root / "overleaf.zip"
        pdf = self.root / "accepted.pdf"
        output = self.root / "latex-output"
        supplemental = self.root / "supplemental.txt"
        linework(image, (2700, 1350))
        latex_fixture(source, image)
        complete_pdf(pdf, image)
        supplemental.write_text("Synthetic supplemental material.\n", encoding="utf-8")

        result = self.run_packager(
            str(source),
            "--paper-id", "JCISE-26-1003",
            "--journal", "jcise",
            "--main", "main.tex",
            "--full-pdf", str(pdf),
            "--output", str(output),
            "--due-date", "2026-09-01",
            "--instruction", "Upload multipart figures separately",
            "--portal-rule", "Corresponding author uploads the complete set",
            "--supplemental-file", str(supplemental),
            "--require-section", "Data Availability",
        )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        self.assertTrue((output / "upload/native/references.bib").is_file())
        root_tex = (output / "upload/native/main.tex").read_text()
        self.assertIn("Production text-only override", root_tex)
        self.assertFalse((output / "upload/native/figure1.png").exists())
        self.assertTrue((output / "upload/figures/Fig_1.tiff").is_file())
        self.assertTrue((output / "upload/supplemental/supplemental.txt").is_file())
        self.assertTrue((output / "JCISE-26-1003_submission_package.zip").is_file())
        manifest = json.loads((output / "manifest.json").read_text())
        self.assertTrue(manifest["manual_review_required"])
        requirements = manifest["controlling_requirements"]
        self.assertEqual(requirements["due_date"], "2026-09-01")
        self.assertEqual(requirements["editorial_instructions"], ["Upload multipart figures separately"])
        self.assertEqual(requirements["supplemental_files"], ["supplemental.txt"])

    def test_nocite_star_blocks_package(self) -> None:
        image = self.root / "figure1.png"
        source = self.root / "overleaf.zip"
        pdf = self.root / "accepted.pdf"
        output = self.root / "nocite-output"
        linework(image, (2700, 1350))
        latex_fixture(source, image, nocite=True)
        complete_pdf(pdf, image)

        result = self.run_packager(
            str(source),
            "--paper-id", "JCISE-26-1004",
            "--main", "main.tex",
            "--full-pdf", str(pdf),
            "--output", str(output),
        )

        self.assertEqual(result.returncode, 2, result.stdout + result.stderr)
        self.assertIn("NOCITE_STAR", (output / "QA_REPORT.md").read_text())
        self.assertFalse((output / "JCISE-26-1004_submission_package.zip").exists())

    def test_latex_infers_relative_width_prunes_archives_and_avoids_reference_false_positive(self) -> None:
        image = self.root / "figure1.png"
        source = self.root / "overleaf-relative.zip"
        pdf = self.root / "accepted-relative.pdf"
        output = self.root / "latex-relative-output"
        linework(image, (3000, 1500))
        complete_pdf_with_reference_traps(pdf, image)

        main = r"""\documentclass{fixture}
\usepackage{graphicx}
\begin{document}
Prior work is cited here~\cite{fixture}.
\begin{figure}
\includegraphics[width=\linewidth]{figure1.png}
\caption{A linework design-space diagram.}
\label{fig:space}
\end{figure}
\section*{Data Availability} Available.
\section*{Conflict of Interest} None.
\bibliographystyle{plain}
\bibliography{references}
\end{document}
"""
        fixture_class = r"""\NeedsTeXFormat{LaTeX2e}
\ProvidesClass{fixture}
\LoadClass[twocolumn]{article}
\RequirePackage[paperwidth=8.5in,left=0.83in,right=0.83in]{geometry}
\setlength\columnsep{0.18in}
"""
        bib = "@article{fixture, author={Author, A.}, title={Fixture}, year={2026}, journal={Fixtures}}\n"
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            (root / "main.tex").write_text(main, encoding="utf-8")
            (root / "fixture.cls").write_text(fixture_class, encoding="utf-8")
            (root / "references.bib").write_text(bib, encoding="utf-8")
            (root / "figure1.png").write_bytes(image.read_bytes())
            (root / "archive").mkdir()
            (root / "archive/old.tex").write_text("\\documentclass{article}\n", encoding="utf-8")
            with zipfile.ZipFile(source, "w", zipfile.ZIP_DEFLATED) as archive:
                for item in sorted(root.rglob("*")):
                    if item.is_file():
                        archive.write(item, item.relative_to(root))

        result = self.run_packager(
            str(source),
            "--paper-id", "JCISE-26-1006",
            "--journal", "jcise",
            "--main", "main.tex",
            "--full-pdf", str(pdf),
            "--output", str(output),
        )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        manifest = json.loads((output / "manifest.json").read_text())
        self.assertGreaterEqual(manifest["figures"][0]["effective_dpi"], 900)
        self.assertTrue((output / "upload/native/fixture.cls").is_file())
        self.assertFalse((output / "upload/native/archive/old.tex").exists())
        report = (output / "QA_REPORT.md").read_text()
        self.assertIn("LATEX_LAYOUT_INFERRED", report)
        self.assertIn("LATEX_UNUSED_FILES_PRUNED", report)
        self.assertNotIn("REFERENCE_NUMBER_GAPS", report)

    def test_blank_complete_pdf_blocks_package(self) -> None:
        image = self.root / "figure1.png"
        source = self.root / "overleaf.zip"
        pdf = self.root / "accepted-with-blank.pdf"
        output = self.root / "blank-pdf-output"
        linework(image, (2700, 1350))
        latex_fixture(source, image)
        complete_pdf_with_blank_page(pdf, image)

        result = self.run_packager(
            str(source),
            "--paper-id", "JCISE-26-1007",
            "--main", "main.tex",
            "--full-pdf", str(pdf),
            "--output", str(output),
        )

        self.assertEqual(result.returncode, 2, result.stdout + result.stderr)
        self.assertIn("BLANK_OR_NEAR_BLANK_PAGE", (output / "QA_REPORT.md").read_text())
        self.assertFalse((output / "JCISE-26-1007_submission_package.zip").exists())


if __name__ == "__main__":
    unittest.main()
